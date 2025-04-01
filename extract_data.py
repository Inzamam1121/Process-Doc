import os
import re
import shutil
import pandas as pd
from docx import Document
import time
import warnings
warnings.filterwarnings("ignore")
from db_data_insert import getDataFromDfandInsertInDB
from datetime import datetime
from logger import general_logger, dir_logger

def initialize_folders(base_directory):
    """Initialize required folders and return their paths."""
    processed_folder = os.path.join(base_directory, "Processed")
    unprocessed_folder = os.path.join(base_directory, "Unprocessed")
    files_folder = os.path.join(unprocessed_folder, "Files")
    links_folder = os.path.join(unprocessed_folder, "Links")

    # Ensure folders exist
    os.makedirs(processed_folder, exist_ok=True)
    os.makedirs(files_folder, exist_ok=True)
    os.makedirs(links_folder, exist_ok=True)

    return processed_folder, unprocessed_folder, files_folder, links_folder

# Regular Expressions with fallback if no table is found in document
patient_regex = re.compile(r'(?:PATIENT|Patient|Name|Patient Information|RE):\s*\n*\s*([\w\s,]+)', re.IGNORECASE)

dob_regex = re.compile(
    r'(?:DOB|Date of Birth|DateOfBirth):\s*\n*\s*(\d{1,4}/\d{1,2}/\d{1,4}|\d{8})'
    r'|Re:\s+.*?\b(\d{1,4}/\d{1,2}/\d{1,4}|\d{8})\b',
    re.IGNORECASE
)
adm_dt_regex = re.compile(r'(?:Date|ADM DT|Admit Date|Request Date|Date ordered|DATE OF CONSULT):\s*\n*\s*(\d{1,4}/\d{1,2}/\d{1,4}(?:\s*\d{1,2}:\d{2}(?:AM|PM)?)?)', re.IGNORECASE)

FOOTER_PATTERN = re.compile(r"([A-Za-z\-']+),\s+([A-Za-z\-'\.]+(?:\s+[A-Za-z\-'\.]+)*)\s+(\d+)?\s+(\d{2}/\d{2}/\d{4})\s+(\d{2}/\d{2}/\d{4})")

def extract_data(doc_path):
    try:
        doc = Document(doc_path)
        patient, dob, adm_dt = None, None, None
        
        # Check if the document has any tables
        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) >= 2:
                        key = row.cells[0].text.strip()
                        value = row.cells[1].text.strip()
                        
                        # Normalize whitespace and remove any trailing colons
                        key_clean = " ".join(key.split()).lower().replace(":", "").strip()
                        value_clean = " ".join(value.split()).strip()

                        if key_clean in ["patient", "patient name", "name"]:
                            patient = value_clean
                        elif key_clean in ["date of birth", "dob"]:
                            dob = value_clean.split(" ")[0]  # Remove extra text like age
                        elif key_clean in ["date", "admit date", "date of visit"]:
                            adm_dt = value_clean
        else:
            # Fallback: extract from paragraphs using regex
            text = "\n".join([p.text for p in doc.paragraphs])
            patient_match = patient_regex.search(text)
            dob_match = dob_regex.search(text)
            adm_dt_match = adm_dt_regex.search(text)
            
            patient = patient_match.group(1).strip() if patient_match else None
            if dob_match:
                dob = dob_match.group(1) or dob_match.group(2)
            else:
                dob = None
            adm_dt = adm_dt_match.group(1) if adm_dt_match else None
        
        # Strip off any time portion from adm_dt
        if adm_dt:
            adm_dt = adm_dt.split(" ")[0]

        # Optional: remove unwanted trailing words from patient name if needed
        for unwanted in ["UNIT", "OtherIdNumber", "DATE OF BIRTH", "DOB", "DOB:", "DOB :", "DOB : ", "DOB :  ", "DOB :   ", "PATIENT:", "DATE", "DATE OF"]:
            if patient and patient.endswith(unwanted):
                patient = patient.rsplit(" ", 1)[0]
                
        return patient, dob, adm_dt
    except Exception as e:
        a=1
        return None, None, None
    
# Function to extract data from text files
def extract_from_text(content):
    """Extract patient details from plain text using regex."""
    try:
        patient_match = patient_regex.search(content)
        dob_match = dob_regex.search(content)
        adm_dt_match = adm_dt_regex.search(content)

        patient = patient_match.group(1).strip() if patient_match else None
        if dob_match:
            dob = dob_match.group(1) or dob_match.group(2)
        else:
            dob = None
        adm_dt = adm_dt_match.group(1) if adm_dt_match else None

        return patient, dob, adm_dt
    except Exception as e:
        a=1
        return None, None, None

def extract_footer_text(doc_path):
    """Extract footer text from a Word document."""
    try:
        doc = Document(doc_path)
        footer_texts = []

        for section in doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                footer_texts.extend([p.text.strip() for p in footer.paragraphs if p.text.strip()])

        return " ".join(footer_texts) if footer_texts else None, ""
    except Exception as e:
        a=1
        return None, e

def extract_from_footer(doc_path):
    """Extracts patient details from the footer."""
    footer_text, error = extract_footer_text(doc_path)
    if not footer_text:
        return None, None, None, error

    # Normalize whitespace to prevent matching issues
    footer_text = " ".join(footer_text.split())

    match = FOOTER_PATTERN.search(footer_text)
    if match:
        last_name, first_middle_name, patient_id, dob, admit_date = match.groups(default="")  
        # Construct full name correctly
        patient_name = f"{first_middle_name} {last_name}".strip()
        return patient_name, dob, admit_date, None

    return None, None, None, None

# Function to generate a unique filename
def get_unique_filename(directory, filename, extension):
    """Ensure the filename is unique by appending a number if needed."""
    counter = 1
    unique_filename = f"{filename}{extension}"
    while os.path.exists(os.path.join(directory, unique_filename)):
        unique_filename = f"{filename}_{counter}{extension}"
        counter += 1
    return unique_filename

# Function to check if a file is a .docx file
def is_docx_file(file_path):
    try:
        doc = Document(file_path)
        return bool(doc.paragraphs)  
    except Exception:
        return False

# Function to check if a file is a text file
def is_txt_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return bool(f.read().strip())  # Ensure it's not empty
    except Exception:
        return False

# Function to check if a file is a .doc file (binary format)
def is_doc_file(file_path):
    try:
        with open(file_path, 'rb') as f:
            header = f.read(8)
        return header[:2] == b'\xd0\xcf'  # Matches the old .doc format
    except Exception:
        return False

def clean_patient_name(name):
    # Remove trailing numbers and whitespace, if any.
    return re.sub(r'\s*\d+$', '', name).strip()

# Function to process text files
def process_text_file(file_path):
    """Extract patient details from text files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read().strip()
            patient, dob, adm_dt = extract_from_text(content)  
            return patient, dob, adm_dt
    except Exception as e:
        a=1
        return None, None, None

# Function to convert date formats
def convert_date_format(date_str):
    """Convert date from mm/dd/yyyy or mm/dd/yy to yyyymmdd format, removing any time component.
       Skip invalid dates like 'DATE'.
    """
    if not date_str or date_str.upper() == "DATE":
        return None  # Skip invalid dates
    
    date_str = date_str.split()[0]  # Remove time if present
    
    if re.match(r"^\d{8}$", date_str):  # Already in yyyymmdd format
        return date_str
    elif re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}$", date_str):  # m/d/yyyy or m/d/yy format
        mm, dd, yy = date_str.split("/")
        if len(yy) == 2:  # Convert 2-digit year to 4-digit
            yy = f"20{yy}" if int(yy) <= 30 else f"19{yy}"  # Assume 2000s for 00-30, else 1900s
        return f"{yy}{mm.zfill(2)}{dd.zfill(2)}"
    
    return None  

# Function to convert yyyymmdd to yyyy-mm-dd format
def format_date_for_csv(yyyymmdd):
    """Convert yyyymmdd to yyyy-mm-dd format for CSV output."""
    if yyyymmdd and len(yyyymmdd) == 8:
        return f"{yyyymmdd[:4]}-{yyyymmdd[4:6]}-{yyyymmdd[6:]}"  # Convert to yyyy-mm-dd
    return None  # Return None for invalid entries

def process_files_in_current_directory(base_directory, files_list):
    csv_data = []
    current_directory = base_directory
    files_list = os.listdir(current_directory)
    
    for file in files_list:
        if file.endswith(".doc") or file.endswith(".docx"):
            file_path = os.path.join(current_directory, file)

            if is_txt_file(file_path):  # If it's actually a text file
                patient, dob, adm_dt = process_text_file(file_path)
            else:  # Process as a doc/docx file
                patient, dob, adm_dt = extract_data(file_path)

            patient = clean_patient_name(patient) if patient else None

            if patient == "Patient:" or patient == "Patient" or patient == "PATIENT" or patient == "PATIENT:" or patient == "Name" or patient == "Name:" or patient == "Patient Information" or patient == "Patient Information:" or patient == "RE:" or patient == "RE":
                patient = None

            if (patient and dob and adm_dt) or (patient and dob):
                dob = convert_date_format(dob)
                adm_dt = convert_date_format(adm_dt)

                if not adm_dt:
                    base_name = f"{patient.lower()}_{dob}"
                else:
                    base_name = f"{patient.lower()}_{dob}_{adm_dt}"

                safe_base_name = re.sub(r'[^\w]+', '_', base_name)
                extension = os.path.splitext(file)[1]

                new_filename = get_unique_filename(processed_folder, safe_base_name, extension)
                new_file_path = os.path.join(processed_folder, new_filename)

                try:
                    shutil.copy2(file_path, new_file_path)
                    csv_data.append([patient, dob, adm_dt, file, new_filename])
                except Exception as e:
                    a=1

            elif not patient or not dob or not adm_dt:
                patient, dob, adm_dt, error = extract_from_footer(file_path)

                if patient and dob and adm_dt:
                    dob = convert_date_format(dob)
                    adm_dt = convert_date_format(adm_dt)

                    base_name = f"{patient.lower()}_{dob}_{adm_dt}"
                    safe_base_name = re.sub(r'[^\w]+', '_', base_name)
                    extension = os.path.splitext(file)[1]

                    new_filename = get_unique_filename(processed_folder, safe_base_name, extension)
                    new_file_path = os.path.join(processed_folder, new_filename)

                    try:
                        shutil.copy2(file_path, new_file_path)
                        csv_data.append([patient, dob, adm_dt, file, new_filename])
                    except Exception as e:
                        a=1
                else:
                    if error and "Package not found" in str(error):
                        shutil.copy2(file_path, os.path.join(unprocessed_folder, "Links", file))
                    else:
                        shutil.copy2(file_path, os.path.join(unprocessed_folder, "Files", file))
            else:
                shutil.copy2(file_path, os.path.join(unprocessed_folder, "Files", file))

    # Build the DataFrame with all the columns
    columns = ["patient_name", "dob", "request_date", "old_document", "new_document"]
    df = pd.DataFrame(csv_data, columns=columns)
    df.insert(0, "id", range(1, len(df) + 1))
    df["is_deleted"] = 0
    df["dob"] = df["dob"].apply(format_date_for_csv)
    df["request_date"] = df["request_date"].apply(format_date_for_csv)
    
    return df

def split_patient_name(name):
    """Splits patient name into first and last name, ensuring each word is capitalized properly."""
    if not isinstance(name, str) or not name.strip():
        return None, None  # Handle empty or invalid cases

    name_parts = name.lower().split()  # Convert everything to lowercase first
    length = len(name_parts)

    if length == 1:
        return name_parts[0].title(), ""  # One word → First Name only

    split_index = length // 2 + (length % 2)  

    first_name = " ".join(name_parts[:split_index]).title()  
    last_name = " ".join(name_parts[split_index:]).title()   
    
    return first_name, last_name

def process_folder(folder_path, files):
    total_files = len([f for f in os.listdir(folder_path) if f.endswith(('.doc', '.docx'))]) 

    print(f"⚙️ Processing folder: {folder_path} with {total_files} files.")

     # Initialize folders dynamically
    global processed_folder, unprocessed_folder, files_folder, links_folder
    processed_folder, unprocessed_folder, files_folder, links_folder = initialize_folders(folder_path)
    
    # Call the function and keep the DataFrame in memory
    df = process_files_in_current_directory(folder_path, files)

    # Apply the function to split names dynamically
    if df.empty:
        df["patient_first_name"] = ""
        df["patient_last_name"] = ""
    else:
        df[['patient_first_name', 'patient_last_name']] = df['patient_name'].apply(
            lambda name: pd.Series(split_patient_name(name))
        )

    # Add old and new document paths
    df['old_document_path'] = df['old_document'].apply(lambda x: os.path.join(folder_path, x))
    df['new_document_path'] = df['new_document'].apply(lambda x: os.path.join(processed_folder, x))

    # Reorder columns to ensure 'patient_first_name' and 'patient_last_name' are in the 2nd and 3rd position
    desired_order = ['id', 'patient_first_name', 'patient_last_name'] + [col for col in df.columns if col not in ['id', 'patient_first_name', 'patient_last_name']]
    df = df[desired_order]

    # Drop the original column if needed
    df.drop(columns=['patient_name'], inplace=True)

    processed_files = len([f for f in os.listdir(processed_folder) if f.endswith(('.doc', '.docx'))])
    link_files = len([f for f in os.listdir(links_folder) if f.endswith(('.doc', '.docx'))])
    
    if(total_files == 0):
        accuracy = 0
    else:
        accuracy = processed_files / (total_files - link_files ) * 100
    
    #print(f"Accuracy of {folder_path} is: {accuracy:.4f}%")

    if len(df) == 0:
        return False,0
    else:
        result = getDataFromDfandInsertInDB(df)

    time.sleep(1)
    return result,processed_files